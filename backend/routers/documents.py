from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from models import Document, User, Tag, DocumentTag
from routers.auth import get_current_user
from database import get_db
from pydantic import BaseModel
from groq import Groq
import os, io
from dotenv import load_dotenv
from datetime import datetime, timezone
import pdfplumber
import docx

load_dotenv()

groq_client = Groq(api_key=os.environ["GROQ_API_KEY"])
GROQ_MODEL = "llama-3.3-70b-versatile"

router = APIRouter()

def _doc_to_dict(doc, tags: list) -> dict:
    return {
        "id": str(doc.id),
        "title": doc.title,
        "raw_text": doc.raw_text,
        "raw_html": doc.raw_html if hasattr(doc, "raw_html") else None,
        "summary": doc.summary,
        "sender": doc.sender if hasattr(doc, "sender") else None,
        "folder_id": str(doc.folder_id) if doc.folder_id else None,
        "status": doc.status,
        "deleted_at": doc.deleted_at,
        "created_at": doc.created_at,
        "tags": tags,
    }

def doc_with_tags(doc, db) -> dict:
    tags = db.query(Tag).join(
        DocumentTag, DocumentTag.tag_id == Tag.id
    ).filter(DocumentTag.document_id == str(doc.id)).all()
    return _doc_to_dict(doc, [{"id": str(t.id), "name": t.name} for t in tags])

def docs_with_tags_batch(docs: list, db) -> list:
    if not docs:
        return []
    doc_ids = [str(d.id) for d in docs]
    rows = db.query(DocumentTag, Tag).join(
        Tag, DocumentTag.tag_id == Tag.id
    ).filter(DocumentTag.document_id.in_(doc_ids)).all()
    tags_map: dict = {did: [] for did in doc_ids}
    for dt, tag in rows:
        tags_map[str(dt.document_id)].append({"id": str(tag.id), "name": tag.name})
    return [_doc_to_dict(doc, tags_map[str(doc.id)]) for doc in docs]

def generate_summary(text: str) -> str:
    if not text or not text.strip():
        return ""

    excerpt = text.strip()[:8000]
    try:
        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            temperature=0.2,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "당신은 문서 관리 도구 Neatly의 AI 요약 어시스턴트입니다.\n"
                        "사용자가 업로드하거나 저장한 문서를 읽고 핵심을 3~4문장으로 요약하세요.\n\n"
                        "규칙:\n"
                        "1. 첫 문장에 문서의 주제와 목적을 명확히 서술하세요.\n"
                        "2. 핵심 수치, 날짜, 금액, 이름, 장소, 기한이 있으면 반드시 포함하세요.\n"
                        "3. 결론, 권고사항, 후속 조치가 있으면 마지막 문장에 담으세요.\n"
                        "4. 회의록·보고서라면 결정 사항과 담당자를 명시하세요.\n"
                        "5. 계약서·법률 문서라면 핵심 조건과 당사자를 포함하세요.\n"
                        "6. 기술 문서라면 주요 기능·변경사항·버전을 강조하세요.\n"
                        "7. '이 문서는', '요약하면' 같은 도입 표현 없이 바로 내용을 서술하세요.\n"
                        "8. 요약문만 출력하세요. 제목, 번호, 마크다운 기호는 쓰지 마세요."
                    )
                },
                {
                    "role": "user",
                    "content": excerpt
                }
            ],
            max_tokens=450,
        )
        return response.choices[0].message.content
    except Exception:
        raise HTTPException(status_code=502, detail="AI 요약 생성에 실패했습니다. 잠시 후 다시 시도해주세요.")

class DocumentCreate(BaseModel):
    title: str
    summary: str
    folder_id: str | None = None

@router.get("/documents")
def get_documents(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    docs = db.query(Document).filter(
        Document.user_id == current_user.id,
        Document.deleted_at == None
    ).all()
    return docs_with_tags_batch(docs, db)

@router.get("/documents/trash")
def get_trash(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    docs = (
        db.query(Document)
        .filter(
            Document.user_id == current_user.id,
            Document.deleted_at != None,
        )
        .order_by(Document.deleted_at.desc())
        .all()
    )
    return docs_with_tags_batch(docs, db)

@router.get("/documents/{document_id}")
def get_document(
    document_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id,
        Document.deleted_at == None
    ).first()

    if not doc:
        raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다.")

    return doc_with_tags(doc, db)

@router.put("/documents/{document_id}")
def update_document(
    document_id: str,
    data: DocumentCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id,
        Document.deleted_at == None
    ).first()

    if not doc:
        raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다.")

    doc.title = data.title
    doc.raw_text = data.summary
    doc.summary = generate_summary(data.summary)
    doc.folder_id = data.folder_id
    db.commit()
    db.refresh(doc)

    return doc_with_tags(doc, db)

@router.patch("/documents/{document_id}/folder")
def move_document_folder(
    document_id: str,
    folder_id: str | None = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id,
        Document.deleted_at == None
    ).first()

    if not doc:
        raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다.")

    doc.folder_id = folder_id
    db.commit()

    return {"message": "폴더가 변경되었습니다."}

@router.delete("/documents/{document_id}")
def delete_document(
    document_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id,
        Document.deleted_at == None
    ).first()

    if not doc:
        raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다.")

    doc.deleted_at = datetime.now(timezone.utc)
    db.commit()

    return {"message": "삭제되었습니다."}

@router.patch("/documents/{document_id}/restore")
def restore_document(
    document_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id,
        Document.deleted_at != None,
    ).first()

    if not doc:
        raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다.")

    doc.deleted_at = None
    db.commit()
    return {"message": "복원되었습니다."}

@router.post("/documents/{document_id}/restore-gmail")
def restore_gmail_email(
    document_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id,
        Document.deleted_at != None,
    ).first()

    if not doc:
        raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다.")

    if not doc.gmail_message_id or not current_user.gmail_access_token:
        doc.deleted_at = None
        db.commit()
        return {"message": "복원되었습니다.", "gmail": False}

    from gmail_utils import make_credentials, refresh_if_expired
    from googleapiclient.discovery import build
    from googleapiclient.errors import HttpError

    creds = make_credentials(current_user)
    refresh_if_expired(creds, current_user, db)
    service = build("gmail", "v1", credentials=creds)

    try:
        service.users().messages().modify(
            userId="me",
            id=doc.gmail_message_id,
            body={"addLabelIds": ["INBOX"], "removeLabelIds": ["TRASH"]},
        ).execute()
    except HttpError as e:
        if e.resp.status == 404:

            doc.deleted_at = None
            db.commit()
            return {"message": "Gmail에서 찾을 수 없어 Neatly 문서로 복원되었습니다.", "gmail": False}
        raise HTTPException(status_code=500, detail="Gmail 복원에 실패했습니다.")

    db.delete(doc)
    db.commit()
    return {"message": "Gmail 받은편지함으로 복원되었습니다.", "gmail": True}

@router.delete("/documents/{document_id}/permanent")
def permanently_delete_document(
    document_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id,
    ).first()

    if not doc:
        raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다.")

    db.delete(doc)
    db.commit()
    return {"message": "영구 삭제되었습니다."}

def extract_text_from_file(filename: str, content: bytes) -> str:

    ext = filename.rsplit(".", 1)[-1].lower()

    if ext == "pdf":

        with pdfplumber.open(io.BytesIO(content)) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)

    elif ext == "docx":

        doc = docx.Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)

    else:

        return content.decode("utf-8", errors="ignore")

@router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    folder_id: str | None = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):

    allowed = {"pdf", "docx", "txt", "md"}
    ext = (file.filename or "").rsplit(".", 1)[-1].lower()
    if ext not in allowed:
        raise HTTPException(
            status_code=400,
            detail="지원하지 않는 파일 형식입니다. (pdf, docx, txt, md 만 가능)"
        )

    if current_user.plan == "FREE":
        today_start = datetime.now(timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0)
        today_count = db.query(Document).filter(
            Document.user_id == current_user.id,
            Document.created_at >= today_start,
            Document.deleted_at == None,
        ).count()
        if today_count >= 3:
            raise HTTPException(
                status_code=403,
                detail="오늘 작성 가능한 문서 수를 초과했습니다. (Free 플랜: 하루 3개)"
            )

    content = await file.read()
    raw_text = extract_text_from_file(file.filename, content).strip()

    if not raw_text:
        raise HTTPException(status_code=422, detail="파일에서 텍스트를 추출할 수 없습니다.")

    title = (file.filename or "업로드 문서").rsplit(".", 1)[0]

    summary = generate_summary(raw_text)

    new_doc = Document(
        user_id=current_user.id,
        folder_id=folder_id,
        title=title,
        raw_text=raw_text,
        summary=summary,
        status="DONE",
    )
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)

    return doc_with_tags(new_doc, db)

@router.post("/documents")
def create_document(
    data: DocumentCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):

    if current_user.plan == "FREE":
        today_start = datetime.now(timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0)
        today_count = db.query(Document).filter(
            Document.user_id == current_user.id,
            Document.created_at >= today_start,
            Document.deleted_at == None,
        ).count()
        if today_count >= 3:
            raise HTTPException(
                status_code=403,
                detail="오늘 작성 가능한 문서 수를 초과했습니다. (Free 플랜: 하루 3개)"
            )

    new_doc = Document(
        user_id=current_user.id,
        folder_id=data.folder_id,
        title=data.title,
        raw_text=data.summary,
        summary=generate_summary(data.summary),
        status="DONE"
    )

    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)

    return doc_with_tags(new_doc, db)
